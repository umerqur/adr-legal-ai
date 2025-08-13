# app.py
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Dict, Optional
import os
from dotenv import load_dotenv
# Load environment variables
load_dotenv()
import io
from datetime import datetime
import uvicorn

# Document processing imports
import PyPDF2
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Azure OpenAI imports
from openai import AzureOpenAI

# Scikit-learn for document search
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# ================== AZURE CONFIGURATION ==================
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")
AZURE_OPENAI_VERSION = os.getenv("AZURE_OPENAI_VERSION", "2025-04-01-preview")
GPT5_DEPLOYMENT_NAME = os.getenv("GPT5_DEPLOYMENT_NAME", "gpt-5-chat")
# ========================================================

app = FastAPI(
    title="ADR Chambers Legal AI Assistant",
    description="Professional legal document analysis powered by Azure AI",
    version="1.0.0"
)

# CORS (safe to allow all for demo/local)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve /static
app.mount("/static", StaticFiles(directory="static"), name="static")

# ---------------- Models ----------------
class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    response: str
    sources: List[str]
    retrieved_chunks: int

class DocumentUploadResponse(BaseModel):
    filename: str
    status: str
    chunks: Optional[int] = None
    content_length: Optional[int] = None
    error: Optional[str] = None

# --------------- Document logic ---------------
class Document:
    def __init__(self, content: str, metadata: Dict):
        self.content = content
        self.metadata = metadata

class TextSplitter:
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 300):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = ["\n\nArticle ", "\n\nSection ", "\n\nClause ", "\n\n", "\n", ". ", " "]
    
    def split_text(self, text: str) -> List[str]:
        chunks = []
        for separator in self.separators:
            if separator in text:
                sections = text.split(separator)
                current_chunk = ""
                for i, section in enumerate(sections):
                    section_text = (separator + section) if i > 0 else section
                    if len(current_chunk) + len(section_text) <= self.chunk_size:
                        current_chunk += section_text
                    else:
                        if current_chunk:
                            chunks.append(current_chunk.strip())
                        current_chunk = section_text
                if current_chunk:
                    chunks.append(current_chunk.strip())
                break
        
        if not chunks:
            for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
                chunk = text[i:i + self.chunk_size]
                if chunk.strip():
                    chunks.append(chunk.strip())
        
        return [c for c in chunks if len(c.strip()) > 50]

class AzureAIClient:
    def __init__(self):
        self.client = None
        if AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_KEY:
            self.client = AzureOpenAI(
                azure_endpoint=AZURE_OPENAI_ENDPOINT,
                api_key=AZURE_OPENAI_KEY,
                api_version=AZURE_OPENAI_VERSION
            )
        self.deployment_name = GPT5_DEPLOYMENT_NAME
        self.system_prompt = (
            "You are an elite legal AI assistant for ADR Chambers, specializing in arbitration, mediation, "
            "and dispute resolution. Always conclude with: "
            "\"This analysis is for informational purposes only and does not constitute legal advice. "
            "Consult with ADR Chambers legal professionals for specific guidance.\""
        )
    
    def generate_response(self, message: str, context: str = "") -> str:
        if not self.client:
            return ("⚠️ Azure is not configured. Set AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_KEY env vars to enable AI responses.\n\n"
                    "This analysis is for informational purposes only and does not constitute legal advice. "
                    "Consult with ADR Chambers legal professionals for specific guidance.")
        try:
            messages = [{"role": "system", "content": self.system_prompt}]
            if context:
                messages.append({"role": "user", "content": f"LEGAL DOCUMENT CONTEXT:\n{context}\n\nQUESTION: {message}"})
            else:
                messages.append({"role": "user", "content": message})
            
            # Simplified approach - just use the basic parameters that work with your model
            response = self.client.chat.completions.create(
                model=self.deployment_name,
                messages=messages,
                max_completion_tokens=4000
            )
            return response.choices[0].message.content
                    
        except Exception as e:
            return f"⚠️ Azure AI Error: {str(e)}"

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = TextSplitter()
        self.vectorizer = TfidfVectorizer(
            max_features=5000, stop_words='english', ngram_range=(1, 3), max_df=0.85, min_df=2
        )
        self.documents: List[Document] = []
        self.document_vectors = None
        self.is_fitted = False
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed")
        try:
            d = docx.Document(io.BytesIO(file_content))
            paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            tables_text = []
            for table in d.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            combined = "\n".join(paragraphs + tables_text)
            if not combined.strip():
                raise Exception("No text content found in Word document")
            return combined.strip()
        except Exception as e:
            raise Exception(f"Word document extraction failed: {str(e)}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                return file_content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise Exception("Could not decode text file")
    
    def process_file(self, filename: str, file_content: bytes) -> Dict:
        try:
            ext = os.path.splitext(filename)[1].lower()
            if ext == '.pdf':
                text_content = self.extract_text_from_pdf(file_content)
            elif ext == '.docx':
                text_content = self.extract_text_from_docx(file_content)
            elif ext == '.txt':
                text_content = self.extract_text_from_txt(file_content)
            else:
                return {'filename': filename, 'status': 'error', 'error': f'Unsupported file type: {ext}'}
            
            if not text_content.strip():
                return {'filename': filename, 'status': 'error', 'error': 'No text content found in file'}
            
            chunks = self.text_splitter.split_text(text_content)
            for i, chunk in enumerate(chunks):
                self.documents.append(Document(
                    content=chunk,
                    metadata={'filename': filename, 'chunk_id': i, 'upload_date': datetime.now().isoformat()}
                ))
            
            if self.documents:
                texts = [d.content for d in self.documents]
                self.document_vectors = self.vectorizer.fit_transform(texts)
                self.is_fitted = True
            
            return {'filename': filename, 'status': 'success', 'chunks': len(chunks), 'content_length': len(text_content)}
        except Exception as e:
            return {'filename': filename, 'status': 'error', 'error': str(e)}
    
    def search_documents(self, query: str, k: int = 15) -> List[tuple]:
        if not self.is_fitted or not self.documents:
            return []
        try:
            qv = self.vectorizer.transform([query])
            sims = cosine_similarity(qv, self.document_vectors).flatten()
            top = sims.argsort()[-k:][::-1]
            results = []
            for idx in top:
                if sims[idx] > 0.001:
                    results.append((self.documents[idx], sims[idx]))
            if not results and self.documents:
                for idx in top[:3]:
                    results.append((self.documents[idx], sims[idx]))
            return results
        except Exception:
            return []
    
    def get_summary(self) -> Dict:
        unique_files = list({d.metadata.get('filename', 'Unknown') for d in self.documents})
        return {'total_chunks': len(self.documents), 'unique_files': unique_files}
    
    def clear_all(self):
        self.documents = []
        self.document_vectors = None
        self.is_fitted = False

# Global instances
ai_client = AzureAIClient()
doc_processor = DocumentProcessor()

# ----------------- HTML -----------------
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>ADR Chambers - AI Legal Assistant</title>
<style>
/* (same styles as before, trimmed for brevity) */
body{font-family:Segoe UI,Tahoma,Geneva,Verdana,sans-serif;background:#f8f9fa;color:#333}
.header{background:linear-gradient(135deg,#8B1538 0%,#A91B47 50%,#C41E3A 100%);color:#fff;padding:2rem 1rem;text-align:center}
.logo{max-width:200px;height:auto;margin-bottom:1rem}
.header h1{font-size:2.5rem;font-weight:700;margin:.5rem 0}
.status-badge{background:rgba(255,255,255,.2);padding:.75rem 1.5rem;border-radius:25px;display:inline-block;margin-top:1rem}
.main-layout{display:flex;max-width:1400px;margin:0 auto;gap:2rem;padding:2rem}
.chat-section{flex:1;background:#fff;border-radius:15px;box-shadow:0 4px 20px rgba(0,0,0,.1);display:flex;flex-direction:column}
.messages-container{flex:1;padding:1.5rem;overflow-y:auto;max-height:60vh;min-height:400px}
.message{margin-bottom:1.5rem;display:flex}
.message.user{justify-content:flex-end}
.message.assistant{justify-content:flex-start}
.message-content{max-width:80%;padding:1rem 1.5rem;border-radius:18px;font-size:.95rem;line-height:1.5}
.message.user .message-content{background:#8B1538;color:#fff;border-bottom-right-radius:4px}
.message.assistant .message-content{background:#f1f3f4;color:#333;border-bottom-left-radius:4px}
.message-sources{margin-top:.75rem;padding:.5rem;background:rgba(139,21,56,.1);border-radius:8px;font-size:.85rem;color:#8B1538;border-left:3px solid #8B1538}
.input-section{padding:1.5rem;border-top:1px solid #e5e5e5;background:#fff}
.input-container{display:flex;gap:1rem;align-items:center}
.message-input{flex:1;padding:1rem 1.5rem;border:2px solid #e5e5e5;border-radius:25px;font-size:1rem;outline:none}
.message-input:focus{border-color:#8B1538}
.send-button{padding:1rem 2rem;background:linear-gradient(135deg,#8B1538 0%,#A91B47 100%);color:#fff;border:none;border-radius:25px;font-weight:600;cursor:pointer}
.sidebar{width:320px;display:flex;flex-direction:column;gap:1.5rem}
.sidebar-section{background:#fff;padding:1.5rem;border-radius:15px;box-shadow:0 4px 20px rgba(0,0,0,.08);border-left:4px solid #8B1538}
.upload-button{width:100%;padding:1rem;background:linear-gradient(135deg,#8B1538 0%,#A91B47 100%);color:#fff;border:none;border-radius:10px;font-weight:600;cursor:pointer;margin-bottom:.5rem}
.upload-help{font-size:.85rem;color:#666;text-align:center}
.metric{flex:1;text-align:center;padding:1rem;background:#f8f9fa;border-radius:10px;border-top:3px solid #8B1538}
.quick-analysis{background:#fff;margin:2rem;padding:2rem;border-radius:15px;box-shadow:0 4px 20px rgba(0,0,0,.1)}
.analysis-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:1rem}
.analysis-button{padding:1rem;background:linear-gradient(135deg,#f8f9fa 0%,#e9ecef 100%);border:2px solid #8B1538;border-radius:10px;font-weight:600;cursor:pointer;color:#8B1538}
.hidden{display:none}
.loading{opacity:.7;font-style:italic}
@media(max-width:768px){.main-layout{flex-direction:column;padding:1rem}.sidebar{width:100%}.analysis-grid{grid-template-columns:1fr}}
</style>
</head>
<body>
    <div class="header">
        <img src="https://adrchambers.com/wp-content/themes/adrc/images/adrc-logo-white.png" alt="ADR Chambers" class="logo">
        <h1>AI Legal Assistant</h1>
        <p>Advanced document analysis for arbitration, mediation, and dispute resolution</p>
        <div class="status-badge"><strong>Legal Disclaimer: This AI provides informational analysis only and does not constitute legal advice. Consult ADR Chambers for professional guidance.</strong></div>
    </div>

    <div class="main-layout">
        <div class="chat-section">
            <div class="messages-container" id="messagesContainer"></div>
            <div class="input-section">
                <div class="input-container">
                    <input type="text" id="messageInput" class="message-input" placeholder="Ask anything about your legal documents...">
                    <button id="sendButton" class="send-button" type="button">Send</button>
                </div>
            </div>
        </div>

        <div class="sidebar">
            <div class="sidebar-section">
                <h3>📁 Document Processing</h3>
                <input type="file" id="fileInput" multiple accept=".pdf,.docx,.txt" style="display:none">
                <button id="uploadButton" class="upload-button" type="button">Upload Documents</button>
                <p class="upload-help">Supports Word documents, PDFs, and text files</p>
            </div>

            <div id="documentLibrary" class="sidebar-section hidden">
                <h3>📚 Document Library</h3>
                <div class="metrics" style="display:flex;gap:1rem;margin-bottom:1rem">
                    <div class="metric">
                        <div class="metric-number" id="chunksCount">0</div>
                        <div class="metric-label">Chunks</div>
                    </div>
                    <div class="metric">
                        <div class="metric-number" id="filesCount">0</div>
                        <div class="metric-label">Files</div>
                    </div>
                </div>
                <div id="filesList"></div>
                <button id="clearButton" class="upload-button" style="background:#dc3545;margin-top:1rem">Clear Library</button>
            </div>
        </div>
    </div>

    <div id="quickAnalysis" class="quick-analysis hidden">
        <h3>⚡ Quick Legal Analysis</h3>
        <p>Click any button to run instant analysis on your documents</p>
        <div class="analysis-grid">
            <button class="analysis-button" onclick="quickAnalysis('contract-summary')">📋 Contract Summary</button>
            <button class="analysis-button" onclick="quickAnalysis('risk-assessment')">⚠️ Risk Assessment</button>
            <button class="analysis-button" onclick="quickAnalysis('dispute-clauses')">🔍 Dispute Clauses</button>
            <button class="analysis-button" onclick="quickAnalysis('critical-dates')">📅 Critical Dates</button>
            <button class="analysis-button" onclick="quickAnalysis('ip-confidentiality')">🛡️ IP & Confidentiality</button>
            <button class="analysis-button" onclick="quickAnalysis('financial-terms')">💰 Financial Terms</button>
            <button class="analysis-button" onclick="quickAnalysis('termination-rights')">🚪 Termination Rights</button>
            <button class="analysis-button" onclick="quickAnalysis('jurisdiction-law')">🌍 Jurisdiction & Law</button>
        </div>
    </div>

    <div class="footer" style="background:#f8f9fa;padding:1rem 2rem;border-top:2px solid #8B1538;text-align:center">
        <div class="disclaimer">
            <p style="color:#8B1538;margin:0;font-size:.9rem;max-width:900px;margin:0 auto">
                <strong>⚖️ Legal Disclaimer: This AI provides informational analysis only and does not constitute legal advice. Consult ADR Chambers for professional guidance.</strong>
            </p>
        </div>
    </div>

    <script src="/static/app.js"></script>
</body>
</html>
"""

# ---------------- Routes ----------------
@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    return HTMLResponse(content=HTML_TEMPLATE)

@app.get("/api/health")
async def health_check():
    azure_configured = bool(AZURE_OPENAI_KEY)
    return {"status": "healthy", "azure_configured": azure_configured, "documents_loaded": len(doc_processor.documents)}

@app.post("/api/upload", response_model=List[DocumentUploadResponse])
async def upload_documents(files: List[UploadFile] = File(...)):
    results = []
    for file in files:
        if file.content_type not in [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain'
        ]:
            results.append(DocumentUploadResponse(filename=file.filename, status="error", error="Unsupported file type"))
            continue
        try:
            content = await file.read()
            result = doc_processor.process_file(file.filename, content)
            results.append(DocumentUploadResponse(
                filename=result['filename'],
                status=result['status'],
                chunks=result.get('chunks'),
                content_length=result.get('content_length'),
                error=result.get('error')
            ))
        except Exception as e:
            results.append(DocumentUploadResponse(filename=file.filename, status="error", error=str(e)))
    return results

@app.post("/api/chat", response_model=ChatResponse)
async def chat_with_documents(request: ChatRequest):
    if len(doc_processor.documents) == 0:
        return ChatResponse(response="Please upload legal documents first to begin analysis.", sources=[], retrieved_chunks=0)
    try:
        search_results = doc_processor.search_documents(request.message, k=15)
        if not search_results:
            return ChatResponse(response="No relevant content found for your query. Try rephrasing or upload more documents.", sources=[], retrieved_chunks=0)

        context_parts = []
        sources = set()
        for i, (doc, score) in enumerate(search_results):
            filename = doc.metadata.get('filename', 'Unknown')
            sources.add(filename)
            context_parts.append(f"--- DOCUMENT SECTION {i+1}: {filename} ---\n{doc.content}\n")
        full_context = "\n".join(context_parts)
        if len(full_context) > 25000:
            full_context = full_context[:25000] + "\n\n[Additional content truncated...]"

        ai_response = ai_client.generate_response(request.message, full_context)
        return ChatResponse(response=ai_response, sources=list(sources), retrieved_chunks=len(search_results))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat error: {str(e)}")

@app.get("/api/documents/summary")
async def get_document_summary():
    return doc_processor.get_summary()

@app.delete("/api/documents")
async def clear_documents():
    doc_processor.clear_all()
    return {"message": "All documents cleared successfully"}

if __name__ == "__main__":
    print("🚀 Starting ADR Chambers Legal AI Assistant…")
    print("📱 Frontend: http://localhost:8000")
    port = int(os.getenv("PORT", 8000))
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=False)





